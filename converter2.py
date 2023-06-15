from docx import Document
from docx.shared import Pt
from copy import deepcopy

# 打开原始的 Word 模板
template_path = 'template.docx'
template_doc = Document(template_path)

# 创建新的 Word 文档
new_doc = Document()

# 复制原始模板到新文档
for element in template_doc.element.body:
    new_doc.element.body.append(deepcopy(element))

# 复制原始模板中的段落样式
for paragraph in template_doc.paragraphs:
    new_paragraph = new_doc.add_paragraph()
    new_paragraph.style = paragraph.style

    # 复制段落中的文本和格式
    for run in paragraph.runs:
        new_run = new_paragraph.add_run(run.text)
        new_run.bold = run.bold
        new_run.italic = run.italic
        new_run.underline = run.underline
        new_run.font.name = run.font.name
        new_run.font.size = run.font.size

# 复制原始模板中的表格
for table in template_doc.tables:
    new_table = new_doc.add_table(rows=table.rows.__len__(), cols=table.columns.__len__())

    # 复制原始表格的样式
    new_table.style = table.style

    # 复制原始表格的内容和样式
    for i, row in enumerate(table.rows):
        new_row = new_table.rows[i]
        new_row.height = row.height
        for j, cell in enumerate(row.cells):
            new_cell = new_table.cell(i, j)
            new_cell.text = cell.text

            # 复制单元格中的段落和格式
            if len(cell.paragraphs) > 0:
                for paragraph in cell.paragraphs:
                    new_paragraph = new_cell.add_paragraph()
                    new_paragraph.style = paragraph.style

                    # 复制段落中的文本和格式
                    for run in paragraph.runs:
                        new_run = new_paragraph.add_run(run.text)
                        new_run.bold = run.bold
                        new_run.italic = run.italic
                        new_run.underline = run.underline
                        if run.font.name is not None:
                            new_run.font.name = run.font.name
                        if run.font.size is not None:
                            new_run.font.size = run.font.size

# 保存新文档
new_doc.save('output.docx')
