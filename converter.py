from openpyxl import load_workbook
from docx import Document
from docx.shared import Pt
from copy import deepcopy
# 文件名称可以修改，但必须放在file-converter项目文件夹下，如果不在文件夹下，需要将单引号中文件名改成相对路径的形式
# 加载Excel文件 input.xlsx
workbook = load_workbook('input.xlsx')
sheet = workbook.active

# 加载Word模板 template.docx
document = Document('template.docx')
document_path = 'template.docx'

# 获取Excel表格数据
data = []
for row in sheet.iter_rows(min_row=2, values_only=True):
    row = [str(cell) for cell in row]
    data.append(row)
    # print(row)

# 获取Excel表头
headers = [cell.value for cell in sheet[1]]
# print(headers)

# 根据管道编号/单线号填充Word模板
target_column = '管道编号/单线号'
row_index = 2
page_index = 0

# 设置字体样式和大小
font_name = '楷体'
font_size = 10.5

for row in data:
    target_value = row[headers.index(target_column)]
    if target_value:
        # 将数据填入Word模板的对应位置
        table = document.tables[page_index]
        current_row = table.rows[row_index]

        for cell_index, value in enumerate(row):
            current_cell = current_row.cells[cell_index]
            current_cell.text = str(value)
            # print(current_cell.text)
            # 创建并修改字体
            run = current_cell.paragraphs[0].runs[0]
            run.font.name = font_name
            run.font.size = Pt(font_size)

        row_index += 1

        # 每页最多填写13行数据
        if row_index >= 15:
            row_index = 0
            page_index += 1
            if page_index >= len(document.tables):
                document.add_page_break()

                document.add_table(rows=13, cols=len(headers))








# 保存填充后的Word文档
document.save('output.docx')
