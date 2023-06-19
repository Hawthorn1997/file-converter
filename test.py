from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.shared import Cm
from copy import deepcopy
# 打开原始的 Word 模板
template_path = 'template.docx'
template_doc = Document(template_path)

# 创建新的 Word 文档
new_doc = Document()

# 复制原始模板到新文档
for element in template_doc.element.body:
    new_doc.element.body.append(deepcopy(element))

# 遍历新文档中的段落并删除段后间距
for paragraph in new_doc.paragraphs:
    paragraph.paragraph_format.space_after = Pt(0)

# 遍历新文档中的表格并删除段后间距
for table in new_doc.tables:
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_after = Pt(0)

# 设置特定行的字体和字体大小
table = new_doc.tables[0]  # 假设表格在文档中的索引为0

row_index = 1  # 设置要更改字体的行索引
target_row = table.rows[row_index]
# 设置行高
row_height_cm = 0.57  # 设置行高0.57厘米
# 将厘米转换为磅
row_height_pt = Cm(row_height_cm).emu
# 设置特定行的行高
target_row.height = row_height_pt

for cell in target_row.cells:
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.font.name = "Arial"
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '等线')
            run.font.size = Pt(10.5)  # 设置字体大小

# 设置特定行的特定列的字体和字体大小
row_index = 0

column_index = 0  # 设置要更改字体的列索引
target_cell = table.cell(row_index, column_index)
for paragraph in target_cell.paragraphs:
    for run in paragraph.runs:
        # run.font.name = "Arial"
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
        run.font.size = Pt(10.5)  # 设置字体大小

column_index = 3  # 设置要更改字体的列索引
target_cell = table.cell(row_index, column_index)
for paragraph in target_cell.paragraphs:
    for run in paragraph.runs:
        run.font.name = "Arial"
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '等线')
        run.font.size = Pt(16)  # 设置字体大小


column_index = 10
target_cell = table.cell(row_index, column_index)
for paragraph in target_cell.paragraphs:
    for run in paragraph.runs:
        run.font.name = "Arial"
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '等线')
        run.font.size = Pt(10.5)  # 设置字体大小


start_row_index = 15  # 起始行索引（包含）
end_row_index = 18  # 结束行索引（不包含）
for row_index, row in enumerate(table.rows):
    if start_row_index <= row_index < end_row_index:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.name = "Arial"
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), '等线')
                    run.font.size = Pt(12)  # 设置字体大小


# 保存新文档
new_doc.save('output.docx')
